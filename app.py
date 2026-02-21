"""
Claude History Browser — Flask web application
Powered by Claude Archive search and storage backends.
"""

import io
import os
import re
import secrets
import threading
import uuid
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, session, flash
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import bridge
from session_manager import SessionManager
from providers import UploadProvider, GCSBucketProvider, MAX_VIEWER_DISK


# --- Markdown-to-docx helpers ---

_INLINE_RE = re.compile(r'\*\*(.+?)\*\*|__(.+?)__|`(.+?)`|\*(.+?)\*|_(.+?)_')


def _add_inline_runs(para, text):
    """Parse inline markdown (bold, italic, code) into docx runs."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        # Plain text before this match
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        if m.group(1) or m.group(2):        # **bold** or __bold__
            run = para.add_run(m.group(1) or m.group(2))
            run.bold = True
        elif m.group(3):                     # `code`
            run = para.add_run(m.group(3))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif m.group(4) or m.group(5):       # *italic* or _italic_
            run = para.add_run(m.group(4) or m.group(5))
            run.italic = True
        pos = m.end()
    # Remaining plain text
    if pos < len(text):
        para.add_run(text[pos:])


def _parse_table_row(line):
    """Split a markdown table row into cell texts, stripping outer pipes."""
    cells = line.split("|")
    # Remove empty strings from leading/trailing pipes
    if cells and not cells[0].strip():
        cells = cells[1:]
    if cells and not cells[-1].strip():
        cells = cells[:-1]
    return [c.strip() for c in cells]


def _is_table_separator(line):
    """Check if a line is a markdown table separator like |---|---|."""
    return bool(re.match(r"^\s*\|?[\s\-:]+(\|[\s\-:]+)+\|?\s*$", line))


def _add_markdown_to_doc(doc, md_text):
    """Convert a markdown string into properly formatted docx elements."""
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            continue

        # Markdown table — detect header row followed by separator
        if "|" in line and (i + 1 < len(lines)) and _is_table_separator(lines[i + 1]):
            header_cells = _parse_table_row(line)
            i += 2  # skip header + separator
            body_rows = []
            while i < len(lines) and "|" in lines[i] and not _is_table_separator(lines[i]):
                body_rows.append(_parse_table_row(lines[i]))
                i += 1
            # Create Word table
            num_cols = len(header_cells)
            table = doc.add_table(rows=1 + len(body_rows), cols=num_cols)
            table.style = "Table Grid"
            # Header row
            for ci, cell_text in enumerate(header_cells):
                if ci < num_cols:
                    cell = table.rows[0].cells[ci]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    run = para.add_run(cell_text)
                    run.bold = True
                    run.font.size = Pt(9)
            # Body rows
            for ri, row_cells in enumerate(body_rows):
                for ci, cell_text in enumerate(row_cells):
                    if ci < num_cols:
                        cell = table.rows[ri + 1].cells[ci]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        _add_inline_runs(para, cell_text)
                        for run in para.runs:
                            run.font.size = Pt(9)
            doc.add_paragraph()  # spacing after table
            continue

        # Horizontal rule (---, ***, ___)
        if re.match(r"^\s*([-*_])\s*\1\s*\1[\s\-*_]*$", line):
            from docx.oxml.ns import qn
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(4)
            sep.paragraph_format.space_after = Pt(4)
            pPr = sep._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single", qn("w:sz"): "4",
                qn("w:space"): "1", qn("w:color"): "CCCCCC",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Heading
        hm = re.match(r"^(#{1,6})\s+(.*)", line)
        if hm:
            level = min(len(hm.group(1)) + 2, 9)  # offset so # → level 3
            doc.add_heading(hm.group(2), level=level)
            i += 1
            continue

        # Bullet list item
        bm = re.match(r"^\s*[-*]\s+(.*)", line)
        if bm:
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(para, bm.group(1))
            i += 1
            continue

        # Numbered list item
        nm = re.match(r"^\s*\d+\.\s+(.*)", line)
        if nm:
            para = doc.add_paragraph(style="List Number")
            _add_inline_runs(para, nm.group(1))
            i += 1
            continue

        # Empty line — skip
        if not line.strip():
            i += 1
            continue

        # Regular paragraph with inline formatting
        para = doc.add_paragraph()
        _add_inline_runs(para, line)
        i += 1


# --- Flask app ---

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", secrets.token_hex(32))
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024  # 200 MB upload limit

# --- Session manager for multi-user viewer support ---
_session_mgr = SessionManager()

# Support custom data directories (e.g. GCS FUSE mount on Cloud Run)
# CLAUDE_DATA_DIR  → raw JSONL sessions (mirrors ~/.claude/)
# CLAUDE_ARCHIVE_DIR → SQLite DB + Tantivy index (mirrors ~/claude-archive/)
_claude_data_dir = os.environ.get("CLAUDE_DATA_DIR")
_claude_archive_dir = os.environ.get("CLAUDE_ARCHIVE_DIR")
if _claude_data_dir or _claude_archive_dir:
    import shutil
    from claude_archive.config import Config, SourceConfig, GeneralConfig, BackendsConfig, BackendConfig

    # SQLite WAL mode requires file locking that GCS FUSE doesn't support.
    # Copy the DB to local disk so SQLite can open it normally.
    _local_archive = "/tmp/claude-archive"
    _remote_archive = _claude_archive_dir or str(Path.home() / "claude-archive")
    if os.path.exists(os.path.join(_remote_archive, "archive.db")):
        os.makedirs(_local_archive, exist_ok=True)
        for _f in ("archive.db",):
            _src = os.path.join(_remote_archive, _f)
            _dst = os.path.join(_local_archive, _f)
            if os.path.exists(_src) and not os.path.exists(_dst):
                shutil.copy2(_src, _dst)

    _config = Config(
        source=SourceConfig(claude_dir=_claude_data_dir or str(Path.home() / ".claude")),
        general=GeneralConfig(archive_dir=_local_archive),
        backends=BackendsConfig(
            tantivy=BackendConfig(enabled=False),  # Tantivy index not copied; use SQLite FTS5
        ),
    )
    _default_archive = bridge.ArchiveBridge(config=_config)
else:
    _default_archive = bridge.ArchiveBridge()


def get_archive() -> bridge.ArchiveBridge:
    """Return the ArchiveBridge for the current request.

    If the request has a viewer_id in the session cookie, return the viewer's
    archive.  Otherwise, return the default (owner's) archive.
    """
    viewer_id = session.get("viewer_id")
    if viewer_id:
        viewer = _session_mgr.get(viewer_id)
        if viewer:
            return viewer.archive
        # Stale cookie — clear it
        session.pop("viewer_id", None)
    return _default_archive


@app.context_processor
def inject_viewer_state():
    """Make `is_viewer` available in all templates."""
    return {"is_viewer": "viewer_id" in session and _session_mgr.get(session["viewer_id"]) is not None}


# --- Background cleanup thread ---

def _cleanup_loop():
    while True:
        threading.Event().wait(300)  # every 5 minutes
        _session_mgr.cleanup_expired()

_cleanup_thread = threading.Thread(target=_cleanup_loop, daemon=True)
_cleanup_thread.start()


@app.template_filter("fmt_ts")
def format_timestamp_filter(ts):
    return bridge.format_timestamp(ts)


@app.template_filter("short_id")
def short_id_filter(uuid_str):
    return uuid_str[:8] if uuid_str else ""


@app.route("/")
def index():
    host_id = request.args.get("host")
    arc = get_archive()
    projects = arc.list_projects(host_id=host_id)
    stats = arc.get_stats(host_id=host_id)
    hosts = arc.list_hosts()
    return render_template("index.html", projects=projects, stats=stats,
                           hosts=hosts, current_host=host_id)


@app.route("/project/<path:project_name>")
def project_view(project_name):
    host_id = request.args.get("host")
    arc = get_archive()
    projects = arc.list_projects(host_id=host_id)
    hosts = arc.list_hosts()
    # Find the selected project
    selected = None
    for proj in projects:
        if proj.name == project_name:
            selected = proj
            break
    if not selected:
        return "Project not found", 404
    stats = arc.get_stats(host_id=host_id)
    return render_template("index.html", projects=projects, selected_project=selected,
                           stats=stats, hosts=hosts, current_host=host_id)


@app.route("/session/<session_id>")
def session_view(session_id):
    host_id = request.args.get("host")
    arc = get_archive()
    projects = arc.list_projects(host_id=host_id)
    hosts = arc.list_hosts()
    meta = arc.get_session_meta(session_id)
    if not meta:
        return "Session not found", 404

    # Get all turns once, then slice for display
    all_turns = arc.get_conversation_turns(meta.session_id)
    total_turns = len(all_turns)
    turns = all_turns[:100]

    # Find the parent project so the sidebar expands
    selected_project = None
    for proj in projects:
        if proj.name == meta.project_name:
            selected_project = proj
            break

    return render_template(
        "session.html",
        projects=projects,
        selected_project=selected_project,
        meta=meta,
        turns=turns,
        total_turns=total_turns,
        subagents=[],  # Subagents are included in main message stream
        session_id=meta.session_id,
        hosts=hosts,
        current_host=host_id,
    )


@app.route("/api/session/<session_id>/turns")
def api_session_turns(session_id):
    offset = request.args.get("offset", 0, type=int)
    limit = request.args.get("limit", 100, type=int)
    turns = get_archive().get_conversation_turns(session_id, offset=offset, limit=limit)
    return jsonify(turns)


@app.route("/api/export/<session_id>", methods=["POST"])
def api_export(session_id):
    """Export selected turns as a Word document."""
    data = request.get_json()
    selected_indices = set(data.get("indices", []))
    include_tools = data.get("include_tools", False)
    include_thinking = data.get("include_thinking", False)

    arc = get_archive()
    all_turns = arc.get_conversation_turns(session_id)
    meta = arc.get_session_meta(session_id)

    doc = Document()

    # Title
    title = doc.add_heading(meta.first_user_message or meta.slug or meta.session_id[:8], level=1)
    doc.add_paragraph(
        f"Project: {meta.project_name}  |  "
        f"Session: {session_id[:8]}  |  "
        f"{bridge.format_timestamp(meta.first_timestamp)} — {bridge.format_timestamp(meta.last_timestamp)}"
    ).style = doc.styles["Subtitle"]

    for i, turn in enumerate(all_turns):
        if i not in selected_indices:
            continue

        if turn["type"] == "user":
            # User heading
            heading = doc.add_heading("You", level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
            ts_para = doc.add_paragraph(bridge.format_timestamp(turn["timestamp"]))
            ts_para.runs[0].font.size = Pt(8)
            ts_para.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            # User text
            doc.add_paragraph(turn["text"])

        elif turn["type"] == "assistant":
            heading = doc.add_heading("Claude", level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0xD4, 0xA5, 0x74)
            ts_para = doc.add_paragraph(bridge.format_timestamp(turn["timestamp"]))
            ts_para.runs[0].font.size = Pt(8)
            ts_para.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            # Text blocks — render markdown formatting
            for text in turn.get("text_blocks", []):
                _add_markdown_to_doc(doc, text)

            # Thinking blocks
            if include_thinking and turn.get("thinking_blocks"):
                doc.add_heading("Thinking", level=3)
                for thought in turn["thinking_blocks"]:
                    _add_markdown_to_doc(doc, thought)

            # Tool calls
            if include_tools and turn.get("tool_calls"):
                doc.add_heading(f"Tool Calls ({len(turn['tool_calls'])})", level=3)
                for tc in turn["tool_calls"]:
                    name = tc.get("tool_name", "")
                    desc = (tc.get("tool_input", {}).get("description")
                            or tc.get("tool_input", {}).get("command", "")[:100]
                            or tc.get("tool_input", {}).get("file_path", "")
                            or tc.get("tool_input", {}).get("pattern", "")
                            or "")
                    para = doc.add_paragraph()
                    run = para.add_run(f"{name}: ")
                    run.bold = True
                    run.font.size = Pt(9)
                    run = para.add_run(desc)
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                    if tc.get("tool_result") and tc["tool_result"].get("text"):
                        result_text = tc["tool_result"]["text"]
                        if len(result_text) > 500:
                            result_text = result_text[:500] + "\n... (truncated)"
                        para = doc.add_paragraph(result_text)
                        for run in para.runs:
                            run.font.size = Pt(8)
                            run.font.name = "Courier New"

        # Separator — thin horizontal line
        sep = doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(6)
        sep.paragraph_format.space_after = Pt(6)
        from docx.oxml.ns import qn
        pPr = sep._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "CCCCCC",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"claude-history-{session_id[:8]}.docx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/search")
def search_view():
    query = request.args.get("q", "").strip()
    fuzzy = request.args.get("fuzzy", "").lower() in ("1", "true", "on")
    host_id = request.args.get("host")
    msg_type = request.args.get("type")  # "user", "assistant", or None (both)
    if msg_type not in ("user", "assistant"):
        msg_type = None
    arc = get_archive()
    results = []
    grouped_results = []
    if query:
        results = arc.search(query, host_id=host_id, fuzzy=fuzzy, message_type=msg_type)
        # Group results by project, preserving overall relevance order
        from collections import OrderedDict
        groups: OrderedDict[str, list] = OrderedDict()
        for r in results:
            key = r.get("project_display") or "Unknown"
            groups.setdefault(key, []).append(r)
        grouped_results = [{"project": k, "hits": v} for k, v in groups.items()]
    projects = arc.list_projects(host_id=host_id)
    hosts = arc.list_hosts()
    stats = arc.get_stats(host_id=host_id)
    return render_template("search.html", projects=projects, query=query,
                           results=results, grouped_results=grouped_results,
                           stats=stats, fuzzy=fuzzy,
                           msg_type=msg_type,
                           hosts=hosts, current_host=host_id)


# --- Multi-user upload / connect routes ---

@app.route("/welcome")
def welcome():
    error = request.args.get("error")
    return render_template("welcome.html", error=error)


@app.route("/upload", methods=["POST"])
def upload():
    file = request.files.get("archive")
    if not file or not file.filename:
        return redirect(url_for("welcome", error="No file selected."))

    # Check disk budget
    if _session_mgr.total_disk_bytes() > MAX_VIEWER_DISK:
        return "Server storage is full. Please try again later.", 503

    viewer_id = str(uuid.uuid4())
    provider = UploadProvider(file)
    result = provider.prepare(viewer_id)

    if result.error:
        return redirect(url_for("welcome", error=result.error))

    _session_mgr.register(result.session)
    session["viewer_id"] = viewer_id
    return redirect(url_for("index"))


@app.route("/connect-gcs", methods=["POST"])
def connect_gcs():
    gcs_path = request.form.get("gcs_path", "").strip()
    if not gcs_path.startswith("gs://"):
        return redirect(url_for("welcome", error="Path must start with gs://"))

    # Check disk budget
    if _session_mgr.total_disk_bytes() > MAX_VIEWER_DISK:
        return "Server storage is full. Please try again later.", 503

    viewer_id = str(uuid.uuid4())
    provider = GCSBucketProvider(gcs_path)
    result = provider.prepare(viewer_id)

    if result.error:
        return redirect(url_for("welcome", error=result.error))

    _session_mgr.register(result.session)
    session["viewer_id"] = viewer_id
    return redirect(url_for("index"))


@app.route("/disconnect", methods=["POST"])
def disconnect():
    viewer_id = session.pop("viewer_id", None)
    if viewer_id:
        _session_mgr.remove(viewer_id)
    return redirect(url_for("welcome"))


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_DEBUG", "1") == "1"
    app.run(debug=debug, host="0.0.0.0", port=port)
